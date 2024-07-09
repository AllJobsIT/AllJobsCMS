import io
import os
import subprocess
from io import BytesIO
from tempfile import NamedTemporaryFile

from babel.dates import format_date
from bs4 import BeautifulSoup
from django.http import FileResponse
from django.shortcuts import get_object_or_404, render
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from docx.shared import Pt, RGBColor

from libs.integer import format_years
from libs.langs import get_language_name
from .models import Worker, WorkExperience


def get_worker_context(worker_id):
    worker = get_object_or_404(Worker, id=worker_id)
    return {
        'worker': worker,
        "grade": ", ".join([grade.value.title for grade in worker.grade]),
        "programming_languages": ", ".join([pl.value for pl in worker.programming_languages]),
        "stack": ", ".join([stack.value for stack in worker.stack]),
        "skills": ", ".join([skills.value for skills in worker.skills]),
        "technologies": ", ".join([technologies.value for technologies in worker.technologies]),
        "databases": ", ".join([databases.value for databases in worker.databases]),
        "development_tools": ", ".join(
            [software_development.value for software_development in worker.software_development]),
        "other_technologies": ", ".join([other_technologies.value for other_technologies in worker.other_technologies]),
        "languages": [{"name": get_language_name(item.value.bound_blocks['language'].value),
                       "grade": item.value.bound_blocks['grade'].value} for item in worker.english_grade],
        "certificates": ", ".join([certificate.value for certificate in worker.certificates]),
        "jobs": WorkExperience.objects.filter(worker_id=worker.id),
        "experience": format_years(int(round(worker.experience))),
    }


def generate_docx(context):
    doc = Document()
    set_document_margins(doc)
    set_default_styles(doc)
    add_heading(doc, context['worker'].full_name())
    add_paragraph(doc, '', context['worker'].get_specialization())
    add_paragraph(doc, 'Грейд:', context.get('grade', '-'))
    add_paragraph(doc, 'Стаж:', context.get('experience', '-'))
    add_skills_and_stack_section(doc, context)
    add_personal_info_section(doc, context)
    add_work_experience_section(doc, context['jobs'])
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


def set_document_margins(doc):
    for section in doc.sections:
        section.top_margin = Pt(36)
        section.bottom_margin = Pt(36)
        section.left_margin = Pt(36)
        section.right_margin = Pt(36)


def set_default_styles(doc):
    styles = doc.styles
    styles.add_style('Citation', WD_STYLE_TYPE.PARAGRAPH)
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)


def add_heading(doc, title, level=1):
    heading = doc.add_heading(level=level)
    heading_run = heading.add_run(title)
    heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    heading_run.font.color.rgb = RGBColor(0, 0, 0)


def add_paragraph(doc, title, text):
    if text != '-':
        paragraph = doc.add_paragraph()
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        title_run = paragraph.add_run(title)
        title_run.bold = True
        title_run.font.color.rgb = RGBColor(102, 0, 0)
        paragraph.add_run(f' {text}')


def add_skills_and_stack_section(doc, context):
    add_heading(doc, 'Навыки и стэк', 2)
    add_table(doc, ["Языки разработки", "Стэк", "Навыки", "Технологии", "Базы данных", "Средства разработки ПО",
                    "Другие технологии"],
              [context.get('programming_languages', '-'), context.get('stack', '-'), context.get('skills', '-'),
               context.get('technologies', '-'), context.get('databases', '-'),
               context.get('software_development', '-'), context.get('other_technologies', '-')])


def add_personal_info_section(doc, context):
    add_heading(doc, 'Личная информация', 2)
    add_language_table(doc, context.get('languages', []))
    add_table(doc, ["Образование", "Сертификаты", "Гражданство", "Место проживания"],
              [context.get('education', '-'), context.get('certificates', '-'),
               context['worker'].citizenship.name or '-', context['worker'].city or '-'])


def add_work_experience_section(doc, jobs):
    add_heading(doc, 'Опыт работы', 2)
    for job in jobs:
        add_heading(doc,
                    f"{job.company_name} ({format_date(job.start_year, format='d MMMM y', locale='ru')} - {format_date(job.end_year, format='d MMMM y', locale='ru')})",
                    3)
        soup = BeautifulSoup(job.description, 'html.parser') if job.description else '-'
        add_table(doc, ["Роль в проекте", "Описание проекта", "Применяемые технологии"],
                  [job.position, soup.get_text() if isinstance(soup, BeautifulSoup) else soup, ", ".join(
                      [other_technologies.value if other_technologies.value else '-' for other_technologies in
                       job.technologies])], True)


def add_table(doc, titles, items, is_work_exp=False):
    table = doc.add_table(rows=0, cols=2)
    table.style = 'Table Grid'
    for i, title in enumerate(titles):
        row_cells = table.add_row().cells
        row_cells[0].text = title
        row_cells[1].text = items[i]
        set_cell_font_color(row_cells[0], RGBColor(102, 0, 0))
        if is_work_exp:
            shade_cell(row_cells[0], "EAD1DC")
            set_cell_font_color(row_cells[0], RGBColor(0, 0, 0))


def add_language_table(doc, languages):
    if languages:
        table = doc.add_table(rows=0, cols=2)
        table.style = 'Table Grid'
        for language in languages:
            row_cells = table.add_row().cells
            row_cells[0].text = language['name']
            row_cells[1].text = language['grade']
            set_cell_font_color(row_cells[0], RGBColor(102, 0, 0))


def set_cell_font_color(cell, color):
    run = cell.paragraphs[0].runs[0]
    run.font.color.rgb = color


def shade_cell(cell, color):
    shading_elm = parse_xml(r'<w:shd {} w:fill="{}"/>'.format(nsdecls('w'), color))
    cell._element.get_or_add_tcPr().append(shading_elm)


def download_docx(request, worker_id):
    context = get_worker_context(worker_id)
    buffer = generate_docx(context)
    worker = get_object_or_404(Worker, id=worker_id)
    response = FileResponse(buffer, as_attachment=True, filename=f"{worker.full_name().replace(' ', '_')}.docx")
    return response


def convert_docx_to_pdf(docx_buffer):
    with NamedTemporaryFile(suffix=".docx", delete=False) as temp_docx_file:
        temp_docx_file.write(docx_buffer.getvalue())
        temp_docx_filename = temp_docx_file.name
    temp_pdf_filename = temp_docx_filename.replace('.docx', '.pdf')
    subprocess.run(['libreoffice', '--headless', '--convert-to', 'pdf', temp_docx_filename, '--outdir',
                    os.path.dirname(temp_docx_filename)])
    with open(temp_pdf_filename, 'rb') as pdf_file:
        pdf_buffer = io.BytesIO(pdf_file.read())
    os.remove(temp_docx_filename)
    os.remove(temp_pdf_filename)
    return pdf_buffer


def download_pdf(request, worker_id):
    context = get_worker_context(worker_id)
    docx_buffer = generate_docx(context)
    pdf_buffer = convert_docx_to_pdf(docx_buffer)
    worker = get_object_or_404(Worker, id=worker_id)
    response = FileResponse(pdf_buffer, as_attachment=True, filename=f"{worker.full_name().replace(' ', '_')}.pdf")
    return response


def view_template(request, worker_id):
    context = get_worker_context(worker_id)
    return render(request, 'html_template_for_worker.html', context=context)
